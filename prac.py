import io
import re
import numpy as np
import cv2
import streamlit as st
import PyPDF2
from PyPDF2 import PdfReader
import pdf2image
from io import BytesIO
from PIL import Image
import openai
import nltk
from nltk.corpus import stopwords
import sklearn.feature_extraction.text as sk 
import time
from transformers import pipeline, set_seed
from st_on_hover_tabs import on_hover_tabs
from docx import Document
import json
from transformers import AutoTokenizer, AutoModelForSequenceClassification

st.set_page_config(layout="wide")

openai.api_key = "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"
RATE_LIMIT={
    "calls":5,
    "period":60
    }
def handle_rate_limit_error():
    # Wait for the rate limit period to elapse
    time.sleep(RATE_LIMIT["period"])

def checker(text):
    while True:
        try:
            
            # Load the zero-shot classification pipeline
            # tokenizer = AutoTokenizer.from_pretrained("facebook/bart-large-mnli")
            model = AutoModelForSequenceClassification.from_pretrained("facebook/bart-large-mnli")
            classifier = pipeline("zero-shot-classification",model=model)
            
            # Set the candidate labels
            labels = ["Human", "AI"]

            # Set the seed for reproducibility
            set_seed(42)

            # Classify the generated text
            result = classifier(text, labels)

            # Print the classification result
            if result['labels'][0] == 'Human':
                st.write(f"The input text is likely to be human-generated with a confidence score of {result['scores'][0]*100:.2f}%")
            else:
                st.write(f"The input text is likely to be AI-generated with a confidence score of {result['scores'][1]*100:.2f}%")
        except openai.error.RateLimitError:
            # Handle rate limit errors by waiting and retrying
            handle_rate_limit_error()


def preprocess(text):
    # Convert text to lowercase
    text = text.lower()
    # Remove numbers
    text = re.sub(r'\d+', '', text)
    # Remove punctuations
    text = re.sub(r'[^\w\s]', '', text)
    # Remove stop words
    nltk.download('stopwords')
    stop_words = set(stopwords.words('english'))
    tokens = nltk.word_tokenize(text)
    filtered_text = [word for word in tokens if word not in stop_words]
    # Join filtered words back into a string
    text = ' '.join(filtered_text)
    return text

def check_plagiarism(text):
    # Preprocess input text
    text = preprocess(text)
    # Initialize TfidfVectorizer
    vectorizer = TfidfVectorizer()
    # Fit and transform input text
    vectors = vectorizer.fit_transform([text])
    # Read text file with reference text
    with open("reference.txt", "r") as f:
        reference_text = f.read()
    # Preprocess reference text
    reference_text = preprocess(reference_text)
    # Transform reference text
    reference_vector = vectorizer.transform([reference_text])
    # Calculate cosine similarity between input text and reference text
    similarity = (vectors * reference_vector.T).toarray()[0][0]
    return similarity

def check_plagiarism_gpt3(text):
    # Set prompt
    prompt = f"Is the following text plagiarized?\n\n{text}\n\nAnswer:"
    # Generate response from GPT-3
    while True:
        try:
            response = openai.Completion.create(
                engine="davinci",
                prompt=prompt,
                max_tokens=1024,
                n=1,
                stop="###",
                temperature=0.5,
            )
            return response.choices[0].text.strip()
        except openai.error.RateLimitError:
            # Handle rate limit errors by waiting and retrying
            handle_rate_limit_error()

def is_blurry(image):
        # Convert the image to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Compute the Laplacian variance
        laplacian = cv2.Laplacian(gray, cv2.CV_64F).var()
        
        # Determine if the image is blurry or not
        if laplacian < 100:
            return True
        else:
            return False

def extract_font_details(run):
    """
    Extracts the font details for a given run in a Word document.
    """
    return {
        "font": run.style.font.name,
        "size": str(run.font.size.pt),
        "style": "bold" if run.bold else "italic" if run.italic else "normal"
    }

def read_docx_file(file):
    doc = Document(file)
    return doc

def main(): 
    st.markdown('<style>' + open('./style.css').read() + '</style>', unsafe_allow_html=True)

    with st.sidebar:
        ch = on_hover_tabs(tabName=['HomePage', 'View',"Image","Check","Format Checker","AI Generated Content Checking"],
                             iconName=['homepage', 'square', 'check','image','check','check'],
                             styles = {'navtab': {'background-color':'#0000FF',
                                                  'color': 'white',
                                                  'font-size': '18px',
                                                  'transition': '.3s',
                                                  'white-space': 'nowrap',
                                                  'text-transform': 'uppercase'},
                                       'tabOptionsStyle': {':hover :hover': {'color': 'Black',
                                                                      'cursor': 'pointer'}},
                                       'iconStyle':{'position':'fixed',
                                                    'left':'7.5px',
                                                    'text-align': 'left'},
                                       'tabStyle' : {'list-style-type': 'none',
                                                     'margin-bottom': '30px',
                                                     'padding-left': '30px'}},
                             default_choice=0)

    if ch=="HomePage":
        st.markdown("<h2 style='font-size:100px'>Paperpal</h2>",unsafe_allow_html=True)
        st.markdown("<h4 style='font-size:20px;color:black'>By researcher.life</h4>",unsafe_allow_html=True)
        st.write("\n")
        st.write("\n")
        st.markdown("<h5 style='font-size:30px;color:black'>To reduce the friction of desk rejection by validating minor and significant checks.</h5>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px;color:black'>To reduce the time and efforts for manually scanning through the manuscript for verifying format and standardization.</h5>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px;color:black'>We @CACTUS have built PayPal</h5>",unsafe_allow_html=True)
        
        st.image("gif.gif", width=900)        
        st.sidebar.title("About")
 
    if ch=="Format Checker":
        st.markdown("<h2 style='font-size:60px'>Format Checker</h2>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px'>Upload a PDF file</h5>",unsafe_allow_html=True)
        pdf_file = st.file_uploader("", type=["docx"])

        if pdf_file is not None:
            doc = read_docx_file(pdf_file)
             # Initialize an empty dictionary to store the formatting information
            formatting_info = {}
            standard={
                "Title":{
                "font":"Arial",
                "size":"12.0",
                "bold":True,
                "italic":False,
                "underline":False
                },
                "Abstract": {
                    "font": "Arial",
                    "size": "12.0",
                    "bold": False,
                    "italic": False,
                    "underline": False
                },
                "Body": {
                    "font": "Calibri",
                    "size": "11.0",
                    "bold": False,
                    "italic": False,
                    "underline": False
                }
            }
            # Get the title formatting
            title = doc.paragraphs[0]
            title_format = {
                'font': title.style.font.name,
                'size': str(title.style.font.size.pt),
                'bold': title.style.font.bold,
                'italic': title.style.font.italic,
                'underline': title.style.font.underline,
            }
            formatting_info['Title'] = title_format

            # Get the abstract formatting
            abstract = doc.paragraphs[1]
            abstract_format = {
                'font': abstract.style.font.name,
                'size': str(abstract.style.font.size.pt),
                'bold': abstract.style.font.bold,
                'italic': abstract.style.font.italic,
                'underline': abstract.style.font.underline,
            }
            formatting_info['Abstract'] = abstract_format

            # Get the body formatting
            body = doc.paragraphs[2]
            body_format = {
                'font': body.style.font.name,
                'size': str(body.style.font.size.pt),
                'bold': body.style.font.bold,
                'italic': body.style.font.italic,
                'underline': body.style.font.underline,
            }
            formatting_info['Body'] = body_format
            
            st.write(formatting_info)
            if standard==formatting_info:
                 st.markdown("<h2 style='font-size:30px'>The formatting is as per the standards.</h2>",unsafe_allow_html=True)
            else:
                 st.markdown("<h2 style='font-size:30px'>THe formatting is not as per the standards.</h2>",unsafe_allow_html=True)
            with open('formatting.json', 'w') as f:
                json.dump(formatting_info, f, indent=4)

    if ch=="AI Generated Content Checking":
        st.markdown("<h2 style='font-size:60px'>AI Generated Content Checker</h2>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px'>Upload a PDF file</h5>",unsafe_allow_html=True)
        pdf_file = st.file_uploader("", type=["pdf"])

        if pdf_file is not None:
            pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(pdf_file.read()))
            text = ""
            images = []
            for page_number in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_number]
                # Extract text from the page
                text += page.extract_text()
        
            checker(text)
        
    if ch=="Image":
        st.markdown("<h2 style='font-size:60px'>Image Quality Checker</h2>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px'>Upload a PDF file</h5>",unsafe_allow_html=True)
        uploaded_file = st.file_uploader("", type=["pdf"])
       # Load the image file into memory
        if uploaded_file is not None:
            images = []
            pdf_reader = PdfReader(io.BytesIO(uploaded_file.read()))
            #one for color scale rgb
            for page_number in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_number]
                Obj = page['/Resources']['/XObject'].getObject()
                for obj in Obj:
                    if Obj[obj]['/Subtype'] == '/Image':
                        # Create image object from binary data
                        img_data = Obj[obj]._data
                        img = Image.open(io.BytesIO(img_data))
                        images.append(img)
                        file_bytes = np.asarray(bytearray(img_data), dtype=np.uint8)
                        image = cv2.imdecode(file_bytes, 1)

                        if is_blurry(image):
                                st.error('The uploaded image is blurry')
                        else:
                                st.success('The uploaded image is not blurry')
                        
                        current_resolution = img.size
                        st.write(current_resolution)
                        desired_resolution = (1000, 1000)
                        scale_factor = max(desired_resolution[0] / current_resolution[0], desired_resolution[1] / current_resolution[1])
                        # Check if increasing the resolution would cause the image to become pixelated or blurry
                        if scale_factor > 2:
                            st.write('Increasing the resolution of this image may cause it to become pixelated or blurry.')
                        else:
                            st.write('Increasing the resolution of this image should not cause it to become pixelated or blurry.')
                        
                        # Display the uploaded image
                        st.image(image, caption='Uploaded Image')
            if len(images) == 0:
                st.write("NO IMAGES FOUND")
                        
    if ch=='View':
        #Pdf Print
        st.markdown("<h2 style='font-size:60px'>View the File</h2>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px'>Upload a PDF file</h5>",unsafe_allow_html=True)
        pdf_uploaded = st.file_uploader("", type="pdf")
        button = st.button("Confirm")
        st.markdown("<h5 style='font-size:30px'>PDF look like</h5>",unsafe_allow_html=True)
        if button and pdf_uploaded is not None:
            if pdf_uploaded.type == "application/pdf":
                images = pdf2image.convert_from_bytes(pdf_uploaded.read())
                for i, page in enumerate(images):
                    st.image(page, use_column_width=True)
                    

    if ch=='Check':
        #Text Retrive
        st.markdown("<h2 style='font-size:60px'>Plagiarism Detector</h2>",unsafe_allow_html=True)
        st.markdown("<h5 style='font-size:30px'>Upload a PDF file</h5>",unsafe_allow_html=True)
        # pdf_uploaded = st.file_uploader("", type="pdf")
        pdf_file = st.file_uploader("", type=["pdf"])

        if pdf_file is not None:
            # Read the PDF file
            pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(pdf_file.read()))
            # pdf_reader = PdfReader(io.BytesIO(pdf_file.read()))
            # Extract text from the PDF
            text = ""
            images = []
            for page_number in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_number]
                # Extract text from the page
                text += page.extract_text()

            # Display the extracted text and images
            user_input=text
            similarity = check_plagiarism(user_input)

            # Check for plagiarism using GPT-3
            answer = check_plagiarism_gpt3(user_input)

            # Display results
            st.subheader("Results:")
            # st.write(f"Similarity score (NLP and IR): {round(similarity, 2)}")
            st.write(f"Plagiarism check (GPT-3): {answer}")
            # if len(text) > 0:
            #     st.write(f"Extracted text from the PDF:")
            #     st.write(text)
            # else:
            #     st.write("No text found in the PDF.")
            


if __name__=='__main__':
    main()